from __future__ import annotations

import os
from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
ASSET_DIR = ROOT / "tmp" / "atomsail-plan-assets"
OUT_FILE = OUT_DIR / "原舟AtomSail项目计划书.docx"
LOGO = ROOT / "frontend" / "web_pc" / "public" / "brand" / "openatom-system-logo-mark.png"
FALLBACK_LOGO = ROOT / "frontend" / "web_pc" / "public" / "logo.png"

FONT_CN = "Hiragino Sans GB"
FONT_LATIN = "Calibri"
NAVY = "0A4B78"
BLUE = "1677FF"
TEAL = "00A6A6"
ORANGE = "FF8A00"
INK = "1D1D1F"
GRAY = "6E6E73"
LIGHT = "F4F6F9"
LIGHT_BLUE = "EAF3FF"
LIGHT_TEAL = "E8F7F5"
BORDER = "D8DEE8"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(run, size=11, bold=None, color=INK, italic=None, font=FONT_CN):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CN)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_specs = {
        "Heading 1": (16, NAVY, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Table Caption" not in styles:
        cap = styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Table Caption"]
    cap.font.name = FONT_LATIN
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    cap.font.size = Pt(9)
    cap.font.color.rgb = rgb(GRAY)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True

    core = doc.core_properties
    core.title = "原舟（AtomSail）项目计划书"
    core.subject = "高校开源社团数字化管理平台项目计划书"
    core.author = "江苏海事职业技术学院开放原子开源社团"
    core.keywords = "原舟, AtomSail, OpenAtom System, 社团管理, 教育数字化, 开源"
    core.comments = "根据 openatom-system 代码仓库与公开权威资料编制。"


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=GRAY)


def add_running_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("原舟 AtomSail  ·  项目计划书")
    set_run_font(r, size=8.5, color=GRAY, bold=True)

    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("江苏海事职业技术学院开放原子开源社团  ·  ")
    set_run_font(r, size=8.5, color=GRAY)
    add_page_number(fp)


def add_para(doc, text, *, bold_lead=None, align=None, color=INK, size=11, after=8, keep=False, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = keep
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=size, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=fill, size="2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_dxa, caption=None, header_fill=LIGHT, font_size=9.5, alignments=None):
    if caption:
        p = doc.add_paragraph(style="Table Caption")
        r = p.add_run(caption)
        set_run_font(r, size=9, color=GRAY, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FBFCFE")
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            if alignments:
                p.alignment = alignments[idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True, color={1: NAVY, 2: BLUE, 3: NAVY}[level])
    return p


def add_toc(doc):
    add_heading(doc, "目录", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Word 中自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, placeholder, end):
        run._r.append(node)
    set_run_font(run, size=10.5, color=GRAY)

    for title in [
        "第一部分  项目概述与执行摘要",
        "第二部分  行业背景与市场分析",
        "第三部分  产品与技术研发（核心章节）",
        "第四部分  竞争分析与核心壁垒",
        "第五部分  市场推广与商业模式",
        "第六部分  团队建设与运营管理",
        "第七部分  财务预测与融资计划",
        "第八部分  社会效益与风险控制",
        "第九部分  发展规划与愿景",
    ]:
        add_para(doc, title, color=NAVY, size=10.5, after=5)


def find_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size, index=1 if bold and "Hiragino" in path else 0)
    return ImageFont.load_default()


def make_financial_chart(path: Path):
    w, h = 1600, 760
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    title_f = find_font(40, True)
    body_f = find_font(26)
    small_f = find_font(22)
    d.text((70, 45), "三年财务预测（单位：万元）", font=title_f, fill=f"#{NAVY}")
    baseline = 640
    top = 135
    max_v = 220
    data = {
        "收入": [18, 72, 210],
        "成本费用": [38, 68, 132],
        "净利润": [-20, 4, 78],
    }
    colors = {"收入": BLUE, "成本费用": TEAL, "净利润": ORANGE}
    years = ["第1年", "第2年", "第3年"]
    for tick in range(0, 221, 50):
        y = baseline - int((tick / max_v) * (baseline - top))
        d.line((130, y, 1510, y), fill="#E6E9EF", width=2)
        d.text((55, y - 14), str(tick), font=small_f, fill="#7A7A7A")
    d.line((130, baseline, 1510, baseline), fill="#AAB2BF", width=3)
    group_x = [340, 800, 1260]
    bar_w = 78
    offsets = [-100, 0, 100]
    for yi, year in enumerate(years):
        d.text((group_x[yi] - 45, 675), year, font=body_f, fill=f"#{INK}")
        for si, name in enumerate(data):
            v = data[name][yi]
            x0 = group_x[yi] + offsets[si] - bar_w // 2
            x1 = x0 + bar_w
            if v >= 0:
                y0 = baseline - int((v / max_v) * (baseline - top))
                y1 = baseline
            else:
                y0 = baseline
                y1 = min(h - 72, baseline + int((abs(v) / max_v) * 320))
            d.rounded_rectangle((x0, y0, x1, y1), radius=12, fill=f"#{colors[name]}")
            d.text((x0 + 12, y0 - 35 if v >= 0 else y1 + 4), str(v), font=small_f, fill=f"#{colors[name]}")
    lx = 980
    for idx, name in enumerate(data):
        x = lx + idx * 185
        d.rounded_rectangle((x, 55, x + 30, 85), radius=6, fill=f"#{colors[name]}")
        d.text((x + 40, 56), name, font=small_f, fill=f"#{INK}")
    im.save(path, quality=95)


def make_market_chart(path: Path):
    w, h = 1600, 760
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    title_f = find_font(40, True)
    body_f = find_font(28, True)
    small_f = find_font(23)
    d.text((70, 45), "目标市场分层与三年切入路径", font=title_f, fill=f"#{NAVY}")
    boxes = [
        ((100, 150, 1500, 330), LIGHT_BLUE, BLUE, "TAM：全国高等学校", "2025年共3,167所；对应社团与学生组织数字化需求的总体机会空间"),
        ((260, 360, 1340, 520), LIGHT_TEAL, TEAL, "SAM：职业本科与高职院校", "共1,641所，组织活动密集、实践育人与社团运营场景契合度高"),
        ((520, 550, 1080, 690), "FFF3E5", ORANGE, "SOM：三年目标", "形成90个付费组织与一批开源社区用户"),
    ]
    for box, fill, accent, title, desc in boxes:
        d.rounded_rectangle(box, radius=28, fill=f"#{fill}", outline=f"#{accent}", width=4)
        d.text((box[0] + 35, box[1] + 24), title, font=body_f, fill=f"#{accent}")
        d.text((box[0] + 35, box[1] + 88), desc, font=small_f, fill=f"#{INK}")
    im.save(path, quality=95)


def cover_page(doc: Document):
    logo = LOGO if LOGO.exists() else FALLBACK_LOGO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(14)
    p.add_run().add_picture(str(logo), width=Inches(1.35))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("原 舟  ATOMSAIL")
    set_run_font(r, size=13, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("项 目 计 划 书")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("高校开源社团数字化管理与协作平台")
    set_run_font(r, size=15, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("让开源协作，有序启航")
    set_run_font(r, size=11.5, color=GRAY, italic=True)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    set_table_borders(table, color=WHITE, size="0")
    metadata = [
        ("项目发起方", "江苏海事职业技术学院开放原子开源社团"),
        ("所属领域", "信息技术服务业 · 教育数字化"),
        ("项目阶段", "工程化建设与校内验证阶段"),
        ("文档版本", "V1.0 · 2026年8月"),
    ]
    for i, (label, value) in enumerate(metadata):
        c1, c2 = table.rows[i].cells
        set_cell_shading(c1, LIGHT)
        set_cell_shading(c2, "FBFCFE")
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p1.add_run(label)
        set_run_font(r, size=9.5, bold=True, color=GRAY)
        p2 = c2.paragraphs[0]
        r = p2.add_run(value)
        set_run_font(r, size=9.5, bold=True if i == 0 else False, color=NAVY if i == 0 else INK)

    add_para(doc, "本计划书中的市场容量、收入、成本与融资数据均为项目规划测算，不构成审计结论或收益承诺。", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=8.5, after=0)
    doc.add_page_break()


def chapter_1(doc):
    add_heading(doc, "第一部分  项目概述与执行摘要", 1)
    add_callout(doc, "一句话定位", "原舟（AtomSail）是一套面向高校开源社团及学生组织的开源化、一体化数字运营平台，将招新、成员、活动、审批、内容、数据与智能文书连接为可追踪的组织工作流。")

    add_heading(doc, "1.1 项目背景与发起初衷", 2)
    add_para(doc, "高校社团承担兴趣培养、实践育人、校园文化建设和学生自治的重要职能，但其日常运营长期依赖微信群、在线表格、纸质申请和个人网盘。人员换届后，账号、文件和流程容易断层；招新、面试、活动、签到、审批等数据分散在不同工具中，既增加干部事务负担，也难以形成可复用的组织知识。开放原子开源社团还需要记录项目贡献、技术方向和社区协作经历，通用办公工具很难完整覆盖这一特殊场景。")
    add_para(doc, "原舟来源于 openatom-system 的持续工程实践。项目以江苏海事职业技术学院开放原子开源社团为首个真实应用场景，以“原”对应开放原子与开源协作，以“舟”呼应海事院校与共同启航，目标是在服务本校社团的同时沉淀一套可复制、可配置、可自主部署的高校社团数字化解决方案。")

    add_heading(doc, "1.2 用户痛点", 2)
    for item in [
        "组织数据分散：成员档案、部门岗位、任职经历、活动记录和荣誉成果缺少统一主数据。",
        "流程依赖人工：招新筛选、面试安排、活动申请、请假与退社审批大量依靠私聊催办。",
        "权限与交接薄弱：账号共享、权限过大、审批责任不清，换届后难以安全移交。",
        "内容与成果难沉淀：活动材料、规章制度、项目成果和成员贡献无法形成长期知识资产。",
        "终端体验割裂：管理端适合复杂配置，但学生更需要在手机端完成报名、签到和消息查看。",
        "智能工具缺少治理：直接使用通用大模型生成材料，难以保证模板一致、数据安全和审核留痕。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 产品方案与核心价值", 2)
    add_para(doc, "原舟以“统一身份—统一组织—统一流程—统一数据—统一触达”为产品主线，形成门户展示、成员服务、运营管理和智能协作四类能力。对学生而言，平台降低了解社团、报名活动、申请加入和查看进度的门槛；对干部而言，平台把重复通知、表格合并和材料整理转化为标准流程；对指导教师与学校管理者而言，平台提供可审计、可追踪的数据依据；对开源社区而言，平台能够成为高校开发者参与真实项目、学习工程协作和贡献代码的实践载体。")

    add_table(doc,
        ["价值对象", "主要价值", "可观测指标"],
        [
            ["普通成员", "统一入口完成报名、申请、签到、消息与成长记录", "任务完成率、活跃率、满意度"],
            ["社团干部", "减少表格汇总和重复通知，形成流程责任链", "流程时长、人工工时、逾期率"],
            ["指导教师/学校", "获得规范、可审计的社团运营数据", "材料完整率、审计留痕率"],
            ["技术团队", "以真实需求驱动开源工程实践和人才培养", "贡献者数、版本频次、缺陷关闭率"],
        ],
        [1500, 4980, 2880],
        caption="表1  原舟的多方价值主张",
    )

    add_heading(doc, "1.4 当前基础与项目阶段", 2)
    add_para(doc, "截至本计划书编制时点，代码仓库已形成较完整的工程基础：包含 Vue 3 Web 门户与管理端、UniApp 小程序、Spring Boot 后端、MySQL 与 Redis 数据层、统一认证与权限体系、文档生成及多类业务模块。仓库快照统计约有49个后端控制器、411个接口映射、80个Vue页面和53个数据库迁移脚本。上述数量仅用于说明工程规模，最终上线范围仍需以测试验收和安全评审结果为准。")
    add_para(doc, "项目阶段定义为“工程化建设与校内验证阶段”：核心业务已具备实现基础，下一阶段重点不是继续无边界增加功能，而是完成主流程收敛、体验一致性、自动化测试、安全合规、可观测性、部署标准化和真实用户验证。")

    add_heading(doc, "1.5 三年目标与关键指标", 2)
    add_table(doc,
        ["维度", "第1年：校内验证", "第2年：区域复制", "第3年：生态扩展"],
        [
            ["产品", "稳定交付Web与移动端核心闭环", "形成多组织SaaS与标准私有化包", "开放插件、API与生态市场"],
            ["用户", "1所学校、10个组织试用", "覆盖3—5个城市、35个付费组织", "累计90个付费组织"],
            ["运营", "核心流程成功率≥98%", "续费率目标≥75%", "续费率目标≥82%"],
            ["技术", "关键接口监控与备份体系落地", "多租户、配置中心与自动化交付", "生态治理与数据智能成熟"],
            ["财务", "完成产品验证，收入18万元", "收入72万元、接近盈亏平衡", "收入210万元、形成正向现金流"],
        ],
        [1200, 2720, 2720, 2720],
        caption="表2  三年阶段目标（规划测算）",
        font_size=9,
    )

    add_heading(doc, "1.6 执行摘要", 2)
    add_para(doc, "原舟选择从“高校开源社团”这一边界清晰、需求真实但供给不足的垂直场景切入。其竞争策略不是与学校大型信息化平台正面替代，而是补足学生组织高频、灵活、跨届的运营层，并通过开放接口与学校现有系统协同。商业上采用“社区版开源 + 云服务订阅 + 私有化部署 + 专业服务”的组合模式；增长上优先依托开放原子社团网络、职业院校合作、技术赛事和校园开源活动形成示范案例；技术上以组织领域模型、流程闭环、安全审计、双端体验和模板化AI能力构成核心壁垒。")


def chapter_2(doc, market_chart: Path):
    doc.add_page_break()
    add_heading(doc, "第二部分  行业背景与市场分析", 1)
    add_heading(doc, "2.1 政策与产业环境", 2)
    add_para(doc, "教育数字化、人工智能赋能教育和开源生态建设共同构成项目发展的宏观基础。教育部等九部门在《关于加快推进教育数字化的意见》中提出坚持应用导向、治理为基，推动教育治理整体性变革，并要求强化网络安全、数据安全、个人信息保护和人工智能安全保障。原舟所处的社团运营场景，具有流程高频、数据分散、组织自治程度高的特点，适合作为教育治理数字化的轻量切入口。")
    add_para(doc, "工业和信息化部关于软件和信息技术服务业的规划强调应用牵引、生态培育、安全可控和开源生态建设。高校开源社团既是开源文化传播节点，也是学生参与真实软件工程的重要组织载体。开放原子开源基金会公开信息显示，截至2025年6月，其已在74所高校建立学生开源社团并举办超过450场开源活动；相关项目已覆盖更广泛的高校和学生群体，证明校园开源组织正在形成可持续网络。")

    add_heading(doc, "2.2 市场规模与结构", 2)
    add_para(doc, "教育部《2025年全国教育事业发展统计公报》显示，全国共有高等学校3,167所，其中普通本科学校1,278所、本科层次职业学校87所、高职（专科）学校1,554所、成人高等学校248所；各种形式的高等教育在学总规模为4,872.57万人。学校数量和学生组织数量并不直接等同于付费市场规模，但为高校社团数字化产品提供了明确的潜在组织基础。")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(market_chart), width=Inches(6.25))
    add_para(doc, "图1  目标市场分层与三年切入路径（学校数量来自教育部2025年统计公报；SOM为项目规划）", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=8.5, after=8)

    add_para(doc, "市场测算采用自下而上的组织订阅法：若将全部3,167所高校视为TAM，并按每校或校内组织平均每年2万元的软件与服务预算估算，理论年机会空间约6.33亿元；若先聚焦87所职业本科和1,554所高职院校，按平均每年1.8万元估算，SAM约2.95亿元。该测算用于战略规划，不代表已有采购预算，也未计入同校多组织、私有化部署和专业服务收入。")

    add_heading(doc, "2.3 目标客户与用户画像", 2)
    add_table(doc,
        ["客户/用户", "核心任务", "主要痛点", "购买或采用动机"],
        [
            ["高校开源社团", "招新、活动、项目与成员成长", "工具分散、交接困难", "低成本获得专业化系统"],
            ["学生会与学生组织", "审批、通知、签到、档案", "流程多、统计压力大", "提升效率与规范性"],
            ["二级学院/团委", "组织监督、资源配置、成果统计", "缺少实时数据和留痕", "提升治理透明度"],
            ["技术社群与实验室", "成员、项目、权限与知识沉淀", "通用OA不适配贡献协作", "连接开源工程实践"],
        ],
        [1700, 2360, 2480, 2820],
        caption="表3  目标客户与用户画像",
        font_size=9,
    )

    add_heading(doc, "2.4 需求趋势判断", 2)
    for item in [
        "从单点工具转向一体化工作流：组织更关注报名—审核—执行—归档的完整闭环。",
        "从PC后台转向双端协同：复杂配置留在Web端，高频服务向小程序和移动端迁移。",
        "从功能采购转向数据治理：权限、审计、数据归属、换届交接和合规成为决策要素。",
        "从通用AI转向场景化AI：用户需要基于组织模板、知识库和审批规则生成可用材料。",
        "从封闭交付转向开放生态：高校技术团队更愿意参与可学习、可二次开发的开源项目。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.5 市场机会与进入窗口", 2)
    add_para(doc, "市场机会来自大型校园平台与碎片化工具之间的空档：大型平台覆盖教务、迎新和统一身份等校级业务，但学生组织的表单变化快、自治程度高、需求迭代频繁；通用协作工具部署快，却缺少成员生命周期、面试、活动、签到、任职和贡献评价等领域模型。原舟可以以轻量接入、开源透明和可配置流程切入，在不替换学校核心系统的前提下提供组织运营层。")


def chapter_3(doc):
    doc.add_page_break()
    add_heading(doc, "第三部分  产品与技术研发（核心章节）", 1)
    add_callout(doc, "研发原则", "先稳定核心闭环，再扩展生态能力；所有智能化功能必须可配置、可追溯、可人工确认，安全与隐私默认优先。", fill=LIGHT_TEAL, accent=TEAL)

    add_heading(doc, "3.1 产品总体架构", 2)
    add_para(doc, "原舟采用“门户与移动服务层—组织运营应用层—业务能力与流程层—数据与基础设施层”的分层架构。门户负责社团品牌、活动、内容、成员成果和合作伙伴展示；小程序承担登录、报名、申请进度、消息和扫码签到；管理后台承担组织、权限、流程、活动、文档和数据配置；后端通过统一认证、领域服务、开放接口和事件机制连接各类终端。")
    add_table(doc,
        ["架构层", "组成", "职责"],
        [
            ["体验层", "Web门户、管理后台、UniApp小程序", "面向访客、成员、干部与管理员提供差异化体验"],
            ["业务层", "成员、招新、活动、审批、签到、通知、内容、文书", "承载社团全生命周期业务与规则"],
            ["平台层", "认证授权、工作流、动态表单、文件、AI、开放API", "提供可复用的横向能力"],
            ["数据层", "MySQL、Redis、对象/文件存储、日志审计", "保障主数据、缓存、文件和可追踪性"],
            ["交付层", "Docker、Nginx、Flyway、CI/CD、备份监控", "支持可重复部署、升级与恢复"],
        ],
        [1400, 3420, 4540],
        caption="表4  原舟产品与技术分层",
    )

    add_heading(doc, "3.2 核心产品模块", 2)
    module_texts = [
        ("组织与成员中心", "统一管理社团、部门、岗位、分组、成员关系、任职记录和状态变更，形成跨届可持续的组织主数据。"),
        ("招新与面试", "支持招新批次、动态报名表、志愿部门、多级审核、面试安排、评分反馈、录取和入社自动建档。"),
        ("活动运营", "覆盖活动创建、报名、名额、签到、通知、获奖记录、复盘与归档，形成活动数据闭环。"),
        ("审批与文书", "支持请假、退社、岗位调整、活动申请等流程，并通过模板生成Word材料，保留版本与审批记录。"),
        ("内容与成员成长", "通过官网、博客、成员主页、项目展示、荣誉和积分沉淀社团成果及个人成长轨迹。"),
        ("统一消息与触达", "整合站内通知、未读提醒、邮件及未来的订阅消息，确保关键节点触达可追踪。"),
        ("开放平台", "提供OAuth/OIDC、OpenAPI、Webhook与第三方应用管理，支持文件、邮箱、实验室等系统协同。"),
        ("智能活动助手", "围绕活动需求澄清、策划案生成、模板字段抽取和正式文档生成提供可审计的AI辅助。"),
    ]
    for title, text in module_texts:
        add_heading(doc, title, 3)
        add_para(doc, text, after=6)

    add_heading(doc, "3.3 关键业务闭环", 2)
    add_number(doc, "招新闭环：发布批次 → 学生申请 → 初审 → 面试 → 终审 → 录取 → 成员建档与权限分配。")
    add_number(doc, "活动闭环：需求提出 → AI/人工策划 → 审批 → 发布报名 → 现场签到 → 成果记录 → 复盘归档。")
    add_number(doc, "成员闭环：加入组织 → 部门岗位 → 活动与项目贡献 → 积分荣誉 → 晋升转岗 → 往届/退社归档。")
    add_number(doc, "治理闭环：权限申请 → 操作执行 → 日志审计 → 异常发现 → 配置整改 → 定期复核。")

    add_heading(doc, "3.4 技术选型与工程实现", 2)
    add_table(doc,
        ["领域", "主要技术", "选择理由"],
        [
            ["前端", "Vue 3、TypeScript、Vite、Element Plus、Tailwind CSS", "组件生态成熟，适合门户与管理后台并行建设"],
            ["移动端", "UniApp、Vue 3", "一套代码支持微信小程序等移动入口"],
            ["后端", "Java 21、Spring Boot 3、MyBatis-Plus", "工程化能力稳定，便于模块化与长期维护"],
            ["认证授权", "Sa-Token、JWT、OAuth 2.0/OIDC、RBAC", "支持统一登录、细粒度权限与第三方接入"],
            ["数据", "MySQL、Redis、Flyway", "适合事务型业务、缓存与版本化数据库升级"],
            ["文档与AI", "Apache POI、模板变量、模型服务代理", "保证正式文书可编辑，并对模型调用集中治理"],
            ["交付", "Docker Compose、Nginx、GitHub Actions", "降低部署门槛并支持持续集成与发布"],
        ],
        [1400, 3480, 4480],
        caption="表5  技术栈与选型依据",
        font_size=9,
    )

    add_heading(doc, "3.5 数据模型与可扩展设计", 2)
    add_para(doc, "数据模型以用户、组织、成员关系、部门岗位和业务事件为主线。成员不是用户表上的单一标签，而是用户与组织之间可记录状态、入会时间、离会时间、部门、岗位和历史的关系实体；分组中心通过统一读取模型兼容部门、签到组、往届分组和外部群组；业务表通过稳定ID关联，避免依赖名称匹配。该设计能够支撑一名用户加入多个组织、跨届保留历史以及未来多租户扩展。")
    add_para(doc, "数据库变更采用Flyway版本化迁移，发布前执行备份和兼容性检查；关键记录原则上采用状态变更与历史留痕，不进行无审计的物理删除。文件与业务记录分离保存，并通过引用、权限和生命周期策略控制访问。")

    add_heading(doc, "3.6 安全、隐私与合规设计", 2)
    for item in [
        "身份安全：密码加密存储，支持会话失效、验证码、限流和关键操作二次确认。",
        "授权安全：RBAC与资源级校验并行，前端隐藏按钮不替代后端权限判断。",
        "数据最小化：公开成员主页使用独立数据视图，手机号、学号、登录凭据和第三方账号标识默认不公开。",
        "审计与追责：记录关键管理操作、审批节点、AI调用摘要、失败原因和操作者。",
        "部署安全：数据库不向公网开放，密钥通过环境变量或加密配置注入，生产环境启用HTTPS、备份和恢复演练。",
        "AI治理：敏感数据脱敏、模型配置集中管理、输出人工确认、提示词与模板版本化，禁止模型直接执行高风险业务操作。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.7 AI活动自动化方案", 2)
    add_para(doc, "智能活动助手不以“自动替代负责人”为目标，而是把活动策划中的重复整理工作结构化。用户输入活动初步想法后，系统先识别目标、对象、时间、地点、预算、人员和风险等缺失字段，再通过多轮问答形成需求确认摘要；用户确认后生成可编辑策划案；策划案再次确认后，将结构化字段映射至学校或社团上传的Word模板，生成活动申请书、活动申请表和志愿者申请表。所有生成内容均保留版本，正式提交前由责任人审核。")
    add_callout(doc, "AI边界", "AI仅提供建议、草稿和字段提取，不自动盖章、不自动提交外部系统、不绕过审批、不替代最终责任人。", fill="FFF3E5", accent=ORANGE)

    add_heading(doc, "3.8 研发计划与质量保障", 2)
    add_table(doc,
        ["阶段", "周期", "研发重点", "退出条件"],
        [
            ["P0 核心稳定", "0—3个月", "认证、组织、招新、活动、报名、签到、消息", "主流程通过验收；P0缺陷清零"],
            ["P1 运营增强", "4—6个月", "审批、文书、积分、成员成长、数据看板", "校内连续运行一个完整活动周期"],
            ["P2 标准化", "7—12个月", "多组织配置、部署脚本、监控、备份、文档", "新组织可在1个工作日内完成初始化"],
            ["P3 生态化", "第2—3年", "插件、开放API、模板市场、跨校协作", "外部贡献者与合作伙伴形成稳定参与"],
        ],
        [1500, 1200, 3900, 2760],
        caption="表6  产品研发路线",
        font_size=9,
    )
    add_para(doc, "质量保障采用需求评审、接口契约、代码审查、自动化测试、预发布验证和上线观察的闭环。关键指标包括核心接口成功率、登录成功率、活动报名成功率、扫码签到成功率、严重缺陷数量、平均恢复时间和备份恢复成功率。首版建议以核心页面崩溃率低于0.5%、主要提交类请求成功率不低于98%作为产品目标，并在真实流量中持续校准。")


def chapter_4(doc):
    doc.add_page_break()
    add_heading(doc, "第四部分  竞争分析与核心壁垒", 1)
    add_heading(doc, "4.1 竞争格局", 2)
    add_para(doc, "原舟面对的不是单一同类产品，而是四类替代方案：微信群与表格等手工工具、钉钉/飞书等通用协作平台、学校现有第二课堂或OA系统、以及针对社团单点需求的小程序或定制系统。不同方案各有优势，但在跨届组织模型、完整业务闭环、开源可控和低成本扩展之间通常需要取舍。")
    add_table(doc,
        ["比较维度", "手工工具", "通用协作平台", "校级平台", "原舟"],
        [
            ["部署门槛", "低", "低", "高", "中，可云端或私有化"],
            ["社团领域模型", "无", "弱", "部分", "强，覆盖成员全周期"],
            ["招新/面试闭环", "人工拼接", "需配置", "视学校而定", "原生支持"],
            ["移动端高频服务", "群聊为主", "较强", "不一致", "Web+小程序协同"],
            ["数据与权限可控", "弱", "平台托管", "较强", "可自部署、可审计"],
            ["开放与二次开发", "无", "受平台限制", "受供应商限制", "开源与开放API"],
            ["AI文书场景", "通用生成", "通用助手", "有限", "模板化、可审核、可追踪"],
        ],
        [1600, 1600, 2000, 1900, 2260],
        caption="表7  替代方案对比",
        font_size=8.8,
    )

    add_heading(doc, "4.2 差异化策略", 2)
    add_para(doc, "原舟不主张替代学校统一身份、教务或财务等核心系统，而是定位为可与其连接的“学生组织运营层”。这一定位可以缩短采购和试点链路，避免大型平台常见的需求排期问题，同时通过标准接口把必要结果回传学校。产品以职业院校和开源社团为首批市场，通过场景深度建立口碑，再逐步向其他学生组织延展。")

    add_heading(doc, "4.3 核心壁垒", 2)
    moat_items = [
        ("真实场景壁垒", "产品由社团自身持续使用和迭代，需求优先级、流程细节和交接痛点来源于真实运营，而非一次性调研。"),
        ("领域数据壁垒", "成员关系、岗位履历、招新评价、活动参与和项目贡献形成连续数据链，长期积累后可支持更准确的组织洞察。"),
        ("工作流与模板壁垒", "将学校材料要求、社团制度和审批规则沉淀为可配置流程与模板，迁移成本随使用时间上升。"),
        ("开源生态壁垒", "代码透明、可自部署和可二次开发有利于技术社团参与贡献，也能降低高校对供应商锁定的顾虑。"),
        ("双端体验壁垒", "Web端处理复杂配置，小程序处理高频轻量操作，统一账号与数据避免重复建设。"),
        ("可信AI壁垒", "AI能力与知识库、模板、字段映射和人工确认结合，比单纯聊天式生成更接近可交付业务结果。"),
    ]
    for title, text in moat_items:
        add_heading(doc, title, 3)
        add_para(doc, text, after=6)

    add_heading(doc, "4.4 壁垒建设计划", 2)
    add_para(doc, "壁垒并非依靠单一专利或功能数量形成，而来自持续积累。第1年重点建立可验证的流程模板、安全基线和真实案例；第2年形成多学校配置经验、实施手册和数据指标体系；第3年通过插件、模板市场、贡献者治理和合作伙伴认证扩大网络效应。与此同时，对品牌名称、Logo、软件著作权、开源许可证和第三方依赖清单进行规范管理，降低知识产权风险。")


def chapter_5(doc):
    doc.add_page_break()
    add_heading(doc, "第五部分  市场推广与商业模式", 1)
    add_heading(doc, "5.1 市场进入策略", 2)
    add_para(doc, "项目采用“校内标杆—同类社团—区域院校—生态伙伴”的渐进式路径。第一阶段以本校开放原子开源社团跑通真实流程，形成可公开演示的案例、数据和部署文档；第二阶段面向开放原子社团网络、江苏省内职业院校和技术类社团开展低成本试点；第三阶段与基金会、开源社区、软件企业、职业教育合作机构共同推广，并形成标准化交付伙伴。")

    add_heading(doc, "5.2 获客渠道", 2)
    for item in [
        "案例传播：发布招新数字化、活动自动化和换届交接的实践报告与演示视频。",
        "校园开源活动：在校源行、技术Meetup、开源大赛和社团交流中提供现场体验。",
        "开发者社区：通过AtomGit/GitHub、技术博客、Issue和贡献者计划吸引技术用户。",
        "院校合作：与团委、二级学院、创新创业学院和信息化部门开展小范围联合试点。",
        "伙伴渠道：为本地软件服务商和学生技术团队提供实施培训与标准化交付包。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.3 产品版本与定价建议", 2)
    add_table(doc,
        ["版本", "目标用户", "建议价格", "主要内容"],
        [
            ["社区版", "技术社团、开发者", "免费开源", "核心组织、成员、活动能力；社区支持"],
            ["云服务基础版", "单个高校社团", "1.2万元/年起", "托管部署、备份、升级、基础支持"],
            ["云服务专业版", "校内多组织", "3.6万元/年起", "多组织、流程模板、数据看板、优先支持"],
            ["私有化部署", "院校或大型组织", "4.8—12万元/次", "专属部署、初始化、培训与验收"],
            ["专业服务", "有定制需求的客户", "按工作量报价", "接口集成、模板配置、迁移与运维"],
        ],
        [1550, 1800, 1750, 4260],
        caption="表8  建议版本与定价（规划口径）",
        font_size=9,
    )
    add_para(doc, "定价应在完成首批访谈和试点后调整。教育场景对持续服务、数据安全和预算周期敏感，合同中应明确服务范围、数据归属、备份恢复、退出机制和定制边界。社区版与商业版的功能边界应透明，避免通过限制数据导出形成不合理锁定。")

    add_heading(doc, "5.4 商业模式", 2)
    add_para(doc, "核心收入由SaaS订阅、私有化部署、实施培训和接口定制构成。长期可增加模板市场、认证培训和生态联合解决方案，但不以学生个人付费或数据变现作为主要商业模式。开源社区版承担品牌传播、产品验证和贡献者增长，商业服务为稳定运维、安全保障和持续研发提供资金，两者形成“开放获客—服务变现—反哺开源”的循环。")

    add_heading(doc, "5.5 销售与客户成功流程", 2)
    add_number(doc, "线索识别：通过活动、内容、社群和伙伴收集意向组织，并完成需求分类。")
    add_number(doc, "轻量诊断：用30—60分钟访谈确认组织规模、现有工具、核心流程和数据安全要求。")
    add_number(doc, "场景演示：只演示与客户痛点相关的2—3个闭环，避免功能堆砌。")
    add_number(doc, "试点验证：选择一个招新批次或活动周期，明确成功指标、责任人和数据边界。")
    add_number(doc, "正式交付：完成配置、迁移、培训、安全检查、验收和运维交接。")
    add_number(doc, "持续成功：按月复盘活跃、流程成功率、问题响应和续费风险，推动模板复用。")

    add_heading(doc, "5.6 品牌传播", 2)
    add_para(doc, "品牌名称“原舟 AtomSail”强调开放原子与海事文化的结合。视觉识别以船锚、原子轨道、协作节点和海浪为核心，传播语为“让开源协作，有序启航”。内容策略围绕真实故事展开，例如“一次招新如何减少多少表格”“活动材料如何从需求到Word申请书”“换届时如何完整移交权限与数据”，以可验证结果替代抽象口号。")


def chapter_6(doc):
    doc.add_page_break()
    add_heading(doc, "第六部分  团队建设与运营管理", 1)
    add_heading(doc, "6.1 组织结构", 2)
    add_para(doc, "项目初期采用精干跨职能团队，核心岗位以学生骨干与指导教师共同治理为主，商业化后逐步补充专职产品、研发、客户成功和安全运维人员。团队不以人数扩张为目标，而以关键职责有人负责、重要决策可追溯、核心知识不集中在单人为基本原则。")
    add_table(doc,
        ["角色", "主要职责", "阶段配置建议"],
        [
            ["项目负责人/产品负责人", "愿景、路线图、需求优先级、外部合作", "1人"],
            ["技术负责人", "架构、安全、代码质量、发布与技术决策", "1人"],
            ["前端与移动端", "Web门户、后台、小程序与体验一致性", "2—3人"],
            ["后端与数据", "领域服务、权限、数据库、开放接口", "2—3人"],
            ["测试与运维", "测试计划、监控、备份、发布和故障响应", "1—2人，可兼岗"],
            ["设计与内容", "品牌、UI规范、文档与案例传播", "1—2人，可兼岗"],
            ["运营与客户成功", "用户培训、社群、试点与续费", "1—2人"],
            ["指导与顾问", "教育场景、合规、安全与商业指导", "按需配置"],
        ],
        [2200, 4960, 2200],
        caption="表9  团队角色与配置建议",
        font_size=9,
    )

    add_heading(doc, "6.2 人才培养与交接机制", 2)
    for item in [
        "建立新成员入门任务：本地运行、修复小缺陷、补测试、完善文档，逐步进入核心模块。",
        "采用双人负责制：关键模块至少一名主责与一名备份，避免毕业或实习导致知识中断。",
        "技术决策文档化：架构变更、数据迁移、权限模型和安全例外必须留下决策记录。",
        "按学期进行权限盘点与密钥轮换，换届时依据清单完成账号、资产和合同移交。",
        "将代码贡献、文档、测试、用户支持和活动组织均纳入贡献评价，避免只认可编码。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.3 研发与运营治理", 2)
    add_para(doc, "需求管理采用季度目标与双周迭代结合：季度明确3—5个可验证结果，双周迭代只承诺有明确验收标准的事项。Issue必须包含背景、用户、范围、验收条件和风险；合并请求至少经过一名非作者评审；涉及权限、财务、隐私和数据迁移的改动必须经过专项复核。上线采用预发布环境验证、数据库备份、分步发布和可回滚方案。")
    add_para(doc, "运营侧建立用户问题分级：影响登录、数据安全或核心流程的问题按P0处理；影响多数用户但有替代方案的问题按P1处理；体验与建议按P2/P3进入迭代池。每月输出产品健康报告，包含活跃组织、流程成功率、异常数量、工单响应、备份情况和安全事件。")

    add_heading(doc, "6.4 社区治理", 2)
    add_para(doc, "项目建议采用明确的开源许可证、贡献指南、行为准则、安全漏洞报告渠道和版本发布规则。维护者依据持续贡献和责任承担产生，重要路线图公开讨论，但涉及个人数据、学校内部配置和安全细节的信息不进入公共仓库。社区版与商业服务共享核心代码，客户专有配置通过扩展机制管理，避免形成难以合并的长期分叉。")


def chapter_7(doc, financial_chart: Path):
    doc.add_page_break()
    add_heading(doc, "第七部分  财务预测与融资计划", 1)
    add_callout(doc, "测算声明", "以下财务数据为基于当前产品阶段、建议定价和三年获客目标的情景测算，未经过审计，不构成盈利保证；正式经营前应结合主体类型、税务政策、采购周期和真实合同重新编制。", fill="FFF3E5", accent=ORANGE)

    add_heading(doc, "7.1 核心测算假设", 2)
    for item in [
        "第1年以校内验证和低价试点为主，形成10个付费组织或等价项目；第2年35个；第3年90个。",
        "收入由云服务订阅、私有化部署和专业服务构成，客单价随客户规模和服务深度提升。",
        "初期团队以学生成员和兼职协作为主，财务模型仍计入合理研发补贴、服务器、测试、安全和运营成本。",
        "不将政府资助、竞赛奖金和一次性捐赠计入主营收入；如获得，将优先用于研发与公益推广。",
        "回款周期按教育客户特点预留，现金流管理应早于利润目标。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.2 三年财务预测", 2)
    add_table(doc,
        ["项目（万元）", "第1年", "第2年", "第3年", "三年合计"],
        [
            ["云服务订阅收入", "8", "36", "126", "170"],
            ["私有化与实施收入", "7", "24", "60", "91"],
            ["培训及专业服务", "3", "12", "24", "39"],
            ["营业收入", "18", "72", "210", "300"],
            ["研发与人员成本", "22", "40", "78", "140"],
            ["服务器与工具", "5", "8", "15", "28"],
            ["市场与客户成功", "4", "10", "22", "36"],
            ["安全、法务及管理", "7", "10", "17", "34"],
            ["成本费用合计", "38", "68", "132", "238"],
            ["净利润（测算）", "-20", "4", "78", "62"],
        ],
        [2600, 1500, 1500, 1500, 2260],
        caption="表10  三年财务预测（规划口径）",
        font_size=9,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(financial_chart), width=Inches(6.25))
    add_para(doc, "图2  三年收入、成本费用与净利润测算", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=8.5, after=8)

    add_heading(doc, "7.3 盈亏平衡与经营指标", 2)
    add_para(doc, "按照上述模型，项目预计在第2年末至第3年初进入稳定盈亏平衡区间。影响盈亏平衡的关键变量不是注册用户数量，而是付费组织数量、续费率、实施交付工时和服务器单位成本。经营看板应重点跟踪年度经常性收入（ARR）、付费组织数、客单价、毛利率、续费率、获客成本、回款周期、工单成本和现金可支撑月数。")

    add_heading(doc, "7.4 融资计划", 2)
    add_para(doc, "建议在完成校内验证并获得2—3个外部试点意向后，申请50万元种子资金或等额的创新创业资助。融资目标不是提前扩大销售团队，而是完成产品标准化、安全合规和可复制交付，使项目从学生工程成果转变为可持续运营的产品。若未获得外部融资，可采用分阶段预算，以订阅和实施回款支持滚动开发。")
    add_table(doc,
        ["资金用途", "比例", "金额（万元）", "主要产出"],
        [
            ["产品与研发", "45%", "22.5", "核心流程、测试、多组织与交付工具"],
            ["安全与合规", "15%", "7.5", "安全评估、隐私制度、备份与灾备演练"],
            ["运营与客户成功", "20%", "10.0", "试点实施、培训、文档与服务体系"],
            ["市场与品牌", "12%", "6.0", "案例内容、活动、演示和伙伴拓展"],
            ["流动资金", "8%", "4.0", "服务器、工具、差旅和应急支出"],
        ],
        [2100, 1200, 1800, 4260],
        caption="表11  50万元种子资金使用计划",
        font_size=9,
    )

    add_heading(doc, "7.5 融资里程碑与投资人回报逻辑", 2)
    add_para(doc, "资金到位后的12个月内，项目应完成：核心P0/P1验收、校内稳定运行、至少3个外部试点、标准部署包、首版安全评估、客户成功手册和可验证的续费意向。投资或资助回报主要来自教育数字化服务收入、可复制的软件资产、开源品牌影响力和高校开发者生态。具体股权、知识产权归属和收益分配需在明确运营主体后依法协商，本计划书不预设股权比例。")


def chapter_8(doc):
    doc.add_page_break()
    add_heading(doc, "第八部分  社会效益与风险控制", 1)
    add_heading(doc, "8.1 社会效益", 2)
    add_para(doc, "原舟首先是一项服务校园治理和开源人才培养的数字化实践。它通过规范社团流程减少干部重复劳动，使成员能够更清晰地了解参与路径；通过记录项目、活动和贡献，帮助学生把零散经历转化为可持续成长档案；通过开放代码和真实需求，为学生提供产品、研发、测试、运维、安全和运营等多角色协作机会；通过可复制的社区版，降低其他高校社团使用数字化工具的门槛。")
    add_table(doc,
        ["效益维度", "预期成果", "建议评价指标"],
        [
            ["治理效率", "流程在线化、责任清晰、材料统一", "流程平均时长、逾期率、材料完整率"],
            ["人才培养", "学生参与真实工程并形成贡献记录", "贡献者数、项目数、成长档案完整率"],
            ["开源文化", "提升开源认知、使用与回馈能力", "活动场次、参与人数、外部贡献数"],
            ["教育公平", "以开源社区版降低中小社团工具成本", "免费部署数、活跃社区数"],
            ["组织可持续", "降低换届造成的数据与知识流失", "交接完成率、文档覆盖率、权限回收率"],
        ],
        [1800, 4180, 3380],
        caption="表12  社会效益评价框架",
        font_size=9,
    )

    add_heading(doc, "8.2 主要风险与应对", 2)
    add_table(doc,
        ["风险", "概率", "影响", "预警信号", "应对措施"],
        [
            ["需求范围失控", "高", "中", "迭代持续延期、P0缺陷积压", "冻结版本范围，按业务闭环排优先级"],
            ["人员流动与毕业交接", "高", "高", "核心模块仅一人掌握", "双人负责、文档化、学期交接演练"],
            ["数据安全与隐私", "中", "高", "越权访问、敏感信息误公开", "最小权限、独立公开视图、审计与安全测试"],
            ["系统稳定性", "中", "高", "报名高峰失败率上升", "压测、限流、监控、备份和故障预案"],
            ["AI输出错误", "中", "中", "模板字段缺失、内容不准确", "结构化校验、人工确认、版本记录"],
            ["采购与回款周期", "中", "中", "试点久不转合同", "分阶段验收、预付款、现金流预警"],
            ["开源许可证与知识产权", "低—中", "高", "依赖来源不清、品牌争议", "SBOM、许可证审查、商标和著作权管理"],
            ["竞争与替代", "中", "中", "通用平台快速补齐功能", "深化领域流程、开放集成和真实案例"],
        ],
        [1500, 850, 850, 2600, 3560],
        caption="表13  风险矩阵与控制措施",
        font_size=8.3,
    )

    add_heading(doc, "8.3 数据与AI合规原则", 2)
    for item in [
        "明确数据控制者、处理者和使用目的，向用户提供清晰的隐私告知与撤回渠道。",
        "仅收集完成业务所必需的数据，设置保存期限、访问范围、导出和删除流程。",
        "个人敏感信息不进入公开页面，不将业务数据用于未经授权的模型训练。",
        "第三方模型调用前进行字段分级与脱敏，记录供应商、模型、用途、耗时和结果状态。",
        "建立未成年人、内容安全、算法输出、数据泄露和账号入侵等事件的应急处置流程。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.4 风险治理机制", 2)
    add_para(doc, "项目建立季度风险评审和重大事件即时上报机制。技术负责人维护安全与稳定性风险，产品负责人维护范围与用户风险，运营负责人维护客户、内容和回款风险，指导教师或顾问对教育场景与合规事项进行复核。高影响风险必须有责任人、截止时间和验证证据，不能仅以“已知悉”关闭。")


def chapter_9(doc):
    doc.add_page_break()
    add_heading(doc, "第九部分  发展规划与愿景", 1)
    add_heading(doc, "9.1 分阶段发展路线", 2)
    add_table(doc,
        ["阶段", "时间", "产品目标", "市场目标", "组织目标"],
        [
            ["启航期", "0—6个月", "核心流程稳定、双端体验统一", "完成本校验证与首批试点", "建立核心维护者与交接机制"],
            ["成长期", "7—18个月", "多组织SaaS、私有化标准包", "形成区域案例与35个付费组织", "补充客户成功和安全运维"],
            ["扩展期", "19—36个月", "开放API、插件和模板生态", "累计90个付费组织", "建立伙伴认证与社区治理"],
            ["生态期", "3年以后", "连接开源人才、项目与组织数据", "形成全国高校协作网络", "多方共建、公益与商业平衡"],
        ],
        [1300, 1350, 2800, 2210, 1700],
        caption="表14  原舟发展路线图",
        font_size=8.8,
    )

    add_heading(doc, "9.2 产品愿景", 2)
    add_para(doc, "原舟的长期目标不是成为功能臃肿的校园超级应用，而是成为高校开源社团可信赖的组织基础设施：成员可以在其中留下成长与贡献，干部可以依靠流程而不是记忆运营组织，学校可以在尊重自治的前提下获得必要治理数据，企业与社区可以更高效地连接真实的高校开发者。")

    add_heading(doc, "9.3 生态愿景", 2)
    add_para(doc, "未来平台将开放流程模板、活动模板、插件接口和数据标准，鼓励不同高校共享经过脱敏和授权的最佳实践。技术社团可以贡献代码和插件，设计与运营成员可以贡献模板、文档和课程，企业与开源社区可以提供真实课题、导师和实践资源。项目通过透明治理保证社区版持续可用，通过商业服务保障专业交付和长期维护。")

    add_heading(doc, "9.4 未来重点方向", 2)
    for item in [
        "贡献导向的人才档案：将代码、文档、测试、活动组织和社区服务形成可验证的多维贡献记录。",
        "跨校开源协作：支持联合活动、项目招募、导师资源和社区课程在多个组织间流转。",
        "可信智能运营：在权限、隐私和人工确认约束下，辅助计划、文书、复盘和知识检索。",
        "开放标准与生态：推动社团组织、活动、成员与贡献数据的可迁移格式，降低平台锁定。",
        "职业教育特色实践：结合海事、软件、智能制造等专业方向，连接校园项目与产业真实需求。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9.5 结语", 2)
    add_para(doc, "原舟起步于一个社团的真实需求，却面向更广泛的高校组织数字化与开源人才培养问题。项目已经具备持续迭代的工程基础，下一步成败取决于能否克制范围、稳定核心流程、守住安全底线、形成真实案例并建立可持续团队。只要坚持“开放、协作、共享、负责”的原则，原舟有望从校内工具成长为连接高校社团、开源社区和产业实践的数字航道。")
    add_callout(doc, "项目愿景", "让每一个高校开源社团都拥有可持续的数字基础，让每一次协作、贡献与成长都被看见。", fill=LIGHT_TEAL, accent=TEAL)


def references(doc):
    doc.add_page_break()
    add_heading(doc, "参考资料与测算说明", 1)
    add_para(doc, "外部市场与政策判断主要参考以下公开资料；产品与技术内容主要依据 openatom-system 代码仓库及仓库内需求、设计和部署文档。", color=GRAY, size=10.5)

    sources = [
        ("教育部：《2025年全国教育事业发展统计公报》，2026年7月6日。", "https://www.moe.gov.cn/jyb_sjzl/sjzl_fztjgb/202607/t20260706_1442870.html"),
        ("教育部等九部门：《关于加快推进教育数字化的意见》（教办〔2025〕3号）。", "https://hudong.moe.gov.cn/srcsite/A01/s7048/202504/t20250416_1187476.html"),
        ("工业和信息化部：《“十四五”软件和信息技术服务业发展规划》解读。", "https://www.miit.gov.cn/jgsj/xxjsfzs/xxfwy/art/2021/art_bfd0f3934e4b40b4a0f4cdb6d747c72d.html"),
        ("开放原子开源基金会：开放原子校源行（太原站）活动及高校开源社团数据。", "https://www.openatom.org/journalism/detail/D36FzC5BJe2J"),
        ("开放原子开源基金会：2024开源春耕计划活动总结。", "https://www.openatom.org/journalism/detail/ayhUni4qK80w"),
    ]
    for idx, (label, url) in enumerate(sources, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(label + " ")
        set_run_font(r, size=9.5, color=INK)
        add_hyperlink(p, url, url)

    add_heading(doc, "测算口径", 2)
    add_para(doc, "TAM/SAM采用高校数量乘以假设年均软件与服务预算的自下而上方法，未将学校数量直接表述为现成客户；SOM、定价、三年收入、成本、利润和融资用途均为规划情景。正式用于创业比赛、项目申报、融资或采购时，应根据参赛规则、实际主体、团队成本、税费、合同和客户访谈结果更新。", size=9.5)


def update_fields_setting(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    financial_chart = ASSET_DIR / "financial-forecast.png"
    market_chart = ASSET_DIR / "market-sizing.png"
    make_financial_chart(financial_chart)
    make_market_chart(market_chart)

    doc = Document()
    configure_document(doc)
    add_running_header_footer(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True
    cover_page(doc)
    add_toc(doc)
    doc.add_page_break()
    chapter_1(doc)
    chapter_2(doc, market_chart)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc, financial_chart)
    chapter_8(doc)
    chapter_9(doc)
    references(doc)
    update_fields_setting(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
